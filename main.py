from docx import Document
from docx.shared import Cm, Pt

doc = Document()
'''Configurar Fonte'''
def conf_font():
    styles = doc.styles['Normal']
    font = styles.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

'''Configurar Margem'''
def conf_margin():
    sections = doc.sections
    top_left = 3.0
    btm_right = 2.5
    for section in sections:
        section.top_margin = Cm(top_left)
        section.left_margin = Cm(top_left)
        section.btm_margin = Cm(btm_right)
        section.right_margin = Cm(btm_right)

'''Escrever as páginas'''
def wrt_page(conteudo):
    for i in range(len(conteudo)):
        
        paginas = conteudo[i]
        for j in range(len(paginas)):
            if(type(paginas).__name__=='str'):
                p = doc.add_paragraph(f'{paginas[j]}')
            else:
        
                textos = paginas[j]
                if(type(textos).__name__=='str'):
                    p = doc.add_paragraph(f'{textos}')
                else:   
                    aux = ''     
                    for k in range(len(textos)):
                        aux += f'{textos[k]}\n' if k<len(textos)-1 else f'{textos[k]}'
                    p = doc.add_paragraph(f'{aux}')
        
        doc.add_page_break()

'''Variaveis'''
estado = 'SP'
faculdade = 'Faculdade de Tecnologia'
tipo_trab = 'Trabalho de Graduação'
cidade = 'Carapicuíba'
ano_letivo = '2022'
curso = 'Analise e Desenvolvimento de Sistemas'
nome = 'Pessoa'
orientador = 'Fulano'
grau = 'Tecnologo'
rnd_text = '\nLorem ipsum dolor sit amet, consectetur adipiscing elit. Praesent vitae ante eget velit consequat viverra. Nunc dui nibh, consectetur a metus ac, accumsan vehicula turpis. Maecenas sit amet eros dignissim, scelerisque ipsum at, viverra nisi. Sed volutpat mauris sit amet est ultrices pulvinar. Proin eu elit est. Suspendisse potenti. Nullam ante nunc, faucibus et tempus in, fermentum et turpis. Etiam maximus et diam rutrum pellentesque. Pellentesque eu scelerisque felis. In elit nunc, laoreet a malesuada vitae, vulputate id leo. Nunc eleifend neque id turpis ultrices, sed sodales turpis finibus. Vivamus rutrum ornare quam, venenatis interdum dui varius at. Mauris tincidunt aliquet eros non euismod. Quisque rutrum metus ac orci vulputate, sit amet facilisis ipsum sagittis.\n'

diretor = 'Dir(a) Beltrano'
banca = 'Banca Examinadora'
spc = ['__________','____________________','______________________________']
data_apr = f'{cidade}, {spc[0]} de {spc[1]} de {spc[0]}'
prf = ['Orientador(a)','Examinador(a)']
ass_prf = [f'Prof(a) Dr(a){spc[2]}\n{prf[0] if i==0 else prf[1]}\n\n' for i in range(3)]

desc_trab = [f'Relatório de Qualificação apresentado junto ao Curso de {curso} da {faculdade} de {cidade} para a obtenção do grau de {grau} na área de {curso}.',f'Orientador(a): {orientador}.']
desc_grad = f'O presente trabalho foi submentido à avaliação da banca examinadora, em cumprimento ás exigências da disciplica de {tipo_trab}, como requisito parcial para a obtenção do grau de {grau} em {curso} de {cidade}'

lst_titulo = ['Lista de Figuras','LIsta de Tabelas','Lista de Abreviaturas, Siglas e Símbolos']

instituto = [f'{faculdade} de {cidade}',f'Curso de {curso}']
grupo = [f'{nome}0{i+1}' for i in range(5)]
cid_ano = [f'{cidade} / {estado}',f'{ano_letivo}']
titulo = ['Titulo','Subtítulo']
frase = ['Basta um dia ruim, para transformar o mais são dos homens, num lunático','Coringa. A Piada Mortal']

secao = [
    '1 Introdução',rnd_text,
    '1.1 Justificativa',rnd_text,
    '1.2 Objetivos',
    '1.2.1 Objetivos do Produto',rnd_text,
    '1.2.2 Objetivos do Projeto',rnd_text,
    '1.3 Metodologia de Pesquisa',rnd_text,
    '1.4 Conceito do Produto',rnd_text,
    '2 Contextualização e Referêncial Teórico',rnd_text,
    '3 Considerações Finais do Projeto e Conclusão',
    '3.1 Qualificação',rnd_text,
    '3.2 Defesa',rnd_text,
    '4 Referências',rnd_text,
    '5 Apêndice',
    '6 Anexo'
]

paginas = [
    [instituto,grupo,titulo,tipo_trab,cid_ano],
    [grupo,titulo,desc_trab,cid_ano],
    [grupo,titulo,desc_grad,data_apr,diretor,banca,ass_prf],
    ['Agradecimentos',rnd_text,frase],
    ['Resumo',rnd_text,'Palavras-Chaves: Palavras.'],
    ['Abstract',rnd_text,'Key-WOrds: Words.'],
    [lst_titulo[0],'Figuras'],
    [lst_titulo[1],'Imagens'],
    [lst_titulo[2],'Siglas'],
    ['Sumário','Indices'],
    [secao]
]

'''
    Executa as Funções e depois salva as modificações do documento
    Execute the functions and then save document's modifications
'''
conf_margin()
conf_font()
wrt_page(paginas)

doc.save('test.docx')